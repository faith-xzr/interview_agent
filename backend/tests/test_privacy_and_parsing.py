from io import BytesIO
from pathlib import Path

from docx import Document

from app.document_parser import extract_text_from_bytes
from app.privacy import mask_pii, restore_pii


def test_mask_pii_replaces_contacts_and_restores_text():
    text = "张三，电话 13812345678，邮箱 zhangsan@example.com，微信 wx_zhangsan。"

    masked = mask_pii(text, candidate_name="张三")

    assert "张三" not in masked.text
    assert "13812345678" not in masked.text
    assert "zhangsan@example.com" not in masked.text
    assert "wx_zhangsan" not in masked.text
    assert "候选人A" in masked.text
    assert "PHONE_1" in masked.text
    assert "EMAIL_1" in masked.text
    assert "WECHAT_1" in masked.text
    assert restore_pii(masked.text, masked.replacements) == text


def test_docx_parser_extracts_document_text():
    document = Document()
    document.add_paragraph("高级 Python 工程师")
    document.add_paragraph("负责 FastAPI、RAG、向量检索系统。")
    buffer = BytesIO()
    document.save(buffer)

    text = extract_text_from_bytes("jd.docx", buffer.getvalue())

    assert "高级 Python 工程师" in text
    assert "FastAPI" in text


def test_plain_text_parser_extracts_utf8_text():
    content = "候选人：李四\n技能：Python、React、SQL".encode("utf-8")

    text = extract_text_from_bytes("resume.txt", content)

    assert "李四" in text
    assert "React" in text


def test_image_pdf_parser_uses_ocr_fallback_for_resume_sample():
    path = Path("samples/resumes/小黄_深度实战版_.pdf")

    text = extract_text_from_bytes(path.name, path.read_bytes())

    assert "小黄" in text
    assert "TikTok" in text
